from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\susan\Documents\grc tool\Control-Centric GRC Tool Implementation Plan.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_run_font(r, bold=True)
        rest = text[len(bold_start):]
        if rest:
            r = p.add_run(rest)
            set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 if level == 0 else 0.45)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    marker = p.add_run("- ")
    set_run_font(marker)
    r = p.add_run(text)
    set_run_font(r)
    return p


def next_numbering_ids(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = []
    num_ids = []
    for node in numbering.findall(qn("w:abstractNum")):
        val = node.get(qn("w:abstractNumId"))
        if val is not None:
            abstract_ids.append(int(val))
    for node in numbering.findall(qn("w:num")):
        val = node.get(qn("w:numId"))
        if val is not None:
            num_ids.append(int(val))
    return (max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1)


def create_decimal_numbering(doc):
    return {"i": 0}


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    num_id["i"] += 1
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    step = p.add_run(f"{num_id['i']}. ")
    set_run_font(step, bold=True)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11),
                     color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(0)
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cells[i].paragraphs[0].add_run(str(item))
            set_run_font(run, size=9.5)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F8F9FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text.strip())
    set_run_font(r, name="Consolas", size=8.5, color=RGBColor(30, 30, 30))


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.text = "Control-Centric GRC Tool Implementation Plan"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Prepared for engineering, product, security, legal, and GRC teams"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)
    return doc


def title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CONTROL-CENTRIC GRC TOOL")
    set_run_font(r, size=23, color=RGBColor(0, 0, 0), bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Implementation Plan for an AI-Native Governance, Risk, and Compliance Platform")
    set_run_font(r, size=14, color=RGBColor(55, 55, 55))

    meta = [
        ("Working name", "Project Ollie"),
        ("Document purpose", "A consolidated, non-repetitive implementation plan suitable for engineering kickoff and cross-functional review"),
        ("Primary design principle", "Controls are the system of record; frameworks, assets, vendors, policies, evidence, risks, AI decisions, and remediation workflows all connect back to controls"),
        ("Initial deployment assumption", "AWS-hosted product experience with Amazon Neptune for graph persistence, S3 Object Lock for immutable evidence, and containerized AI orchestration"),
        ("Planning exclusions", "No dates or budget assumptions are included"),
    ]
    add_table(doc, ["Field", "Decision"], meta, [1900, 7460], header_fill=WHITE)
    add_callout(
        doc,
        "Core Thesis",
        "Traditional GRC platforms often begin with audit checklists, documents, or isolated risk registers. This product should begin with the control graph. Every material answer the tool gives should be explainable by traversing from a control to the assets, evidence, framework requirements, owners, risk scenarios, vendors, and AI decisions connected to it.",
    )
    doc.add_page_break()


def market_context(doc):
    add_heading(doc, "1. Product Intent and Market Baseline", 1)
    add_para(
        doc,
        "The platform should compete in the same expectation space as modern compliance automation, integrated risk management, and third-party risk products, while differentiating through a control-centric graph model and auditable AI workflows. Current market signals show that buyers expect centralized controls and ownership, continuous monitoring, reusable evidence, framework cross-mapping, risk dashboards, remediation workflows, and reporting for executives and auditors."
    )
    add_table(
        doc,
        ["Market expectation", "Implication for this product"],
        [
            ("Controls, ownership, and evidence in one place", "The Control Center must be the first-class workspace. Evidence reuse, owner accountability, and framework mapping should not live in separate modules."),
            ("Automated continuous control monitoring", "The ingestion layer must run as a continuous loop, not a quarterly audit-prep import."),
            ("Integrated risk and control assurance", "Risk, vulnerability, audit findings, regulatory obligations, controls, and assets need a shared data foundation."),
            ("Unified governance, risk, and operational view", "Executive reporting must roll up control health, compliance readiness, operational exposure, and FAIR risk metrics."),
            ("One control library mapped to many frameworks", "Framework import and cross-mapping should be part of the core data model, with human approval for AI-proposed mappings."),
            ("Third-party risk intelligence and ongoing monitoring", "Vendor nodes should connect to controls, assets, data, external assessments, and continuous threat intelligence."),
        ],
        [2500, 6860],
    )
    add_para(
        doc,
        "Design target: build a serious enterprise tool, not a dashboard demo. The product must provide fast time-to-value, but its underlying architecture must also withstand auditor review, legal scrutiny, AI governance review, data-isolation testing, and operational failure modes."
    )


def charter_scope(doc):
    add_heading(doc, "2. Business Charter and Scope", 1)
    add_heading(doc, "2.1 Product Mission", 2)
    add_para(
        doc,
        "Create an AI-native, control-centric GRC platform that continuously measures control health, maps controls to frameworks and enterprise entities, quantifies control-driven risk, and turns evidence collection, audit response, and remediation into governed workflows."
    )
    add_heading(doc, "2.2 Success Outcomes", 2)
    outcomes = [
        "Reduce manual audit preparation by using automated evidence collection, evidence reuse, and control-to-framework cross-mapping.",
        "Give control owners a clear, action-oriented workflow for evidence requests, degraded controls, and remediation tasks.",
        "Give GRC analysts one place to understand whether a control is implemented, where it operates, what evidence proves it, which framework requirements it satisfies, and what risk exposure it affects.",
        "Give executives quantitative risk views, including FAIR-based loss ranges, top risk scenarios, and material control failures.",
        "Give auditors immutable evidence, AI decision lineage, and human approval records tied directly to the relevant controls.",
    ]
    for item in outcomes:
        add_bullet(doc, item)
    add_heading(doc, "2.3 Initial Scope Boundary", 2)
    add_table(
        doc,
        ["Area", "In Scope for MVP", "Deferred"],
        [
            ("Control Center", "Manual control creation, owner assignment, implementation status, KPIs/KRIs, control health history", "Complex control inheritance and advanced enterprise taxonomy tooling"),
            ("Frameworks", "One primary framework, recommended NIST CSF 2.0, plus internal control library", "Full multi-framework import library on day one"),
            ("Assets", "AWS cloud resources, central IAM assets, critical applications, and Tier 1 vendors", "On-prem legacy systems and long-tail asset classes"),
            ("Evidence", "API evidence from first integrations plus manual upload and immutable storage", "Broad connector marketplace and OCR-heavy evidence normalization"),
            ("AI", "Evidence validation, framework mapping recommendations, KMS read-only query assistant, FAIR narrative assistant", "Fully autonomous remediation on mission-critical systems"),
            ("Risk", "FAIR scenario model, Monte Carlo microservice, percentile outputs, risk appetite alerts", "Enterprise stress testing and capital planning integrations"),
        ],
        [1600, 3980, 3780],
    )
    add_heading(doc, "2.4 Non-Goals", 2)
    for item in [
        "Do not build a generic document repository. Documents matter only when connected to controls, evidence, policies, audits, vendors, assets, or risks.",
        "Do not allow AI agents to directly mutate the production graph. All writes must pass through middleware validation and guardrails.",
        "Do not treat FAIR as a simple high/medium/low score. The risk engine should preserve uncertainty, ranges, and scenario context.",
        "Do not make the graph visualization the only user interface. Executives, control owners, and analysts need role-specific workflows.",
    ]:
        add_bullet(doc, item)


def architecture(doc):
    add_heading(doc, "3. Target Architecture", 1)
    add_para(
        doc,
        "The platform should be built around a property graph persisted in Amazon Neptune. The graph model is appropriate because the product's value comes from traversing many-to-many relationships: controls to assets, controls to requirements, controls to evidence, controls to risks, controls to owners, controls to policies, and controls to vendor dependencies."
    )
    add_heading(doc, "3.1 Logical Layers", 2)
    add_table(
        doc,
        ["Layer", "Responsibilities", "Primary Services / Components"],
        [
            ("Product UI", "Role-specific experiences for dashboards, control detail, graph exploration, approval queues, evidence review, policies, vendor risk, and settings", "React on AWS, Cognito or enterprise SSO, Cytoscape.js or similar graph visualization"),
            ("API Gateway and Middleware", "RBAC enforcement, tenant scoping, schema validation, idempotency, HMAC webhook verification, AI action allow-lists, graph query construction", "AWS API Gateway or ALB, containerized service, validation library, policy engine"),
            ("Control Graph", "System of record for controls and relationships", "Amazon Neptune using property graph semantics with Gremlin or openCypher"),
            ("Evidence Store", "Immutable raw evidence, policy artifacts, audit packages, AI context logs, and versioned attachments", "Amazon S3 with Object Lock, versioning, KMS encryption, retention policies"),
            ("AI Orchestration", "Specialized agents for evidence validation, mapping, risk, KMS, vendor assessment, policy drafting, and remediation recommendations", "Containerized orchestrator, model endpoint, structured output, telemetry harness"),
            ("Vector and RAG Index", "Semantic retrieval of potentially relevant controls, policies, evidence, and framework requirements", "Vector index keyed by graph identifiers; graph remains source of truth"),
            ("Risk Engine", "FAIR scenario parameterization, Monte Carlo simulation, percentile reporting, appetite thresholds", "Python microservice with numpy/scipy or equivalent statistical stack"),
            ("Workflow Integrations", "Tickets, notifications, stakeholder tasks, source system sync, legacy import", "Jira, ServiceNow, Slack/Teams, AWS Security Hub, Okta/Entra, vulnerability scanners"),
        ],
        [1900, 4300, 3160],
    )
    add_heading(doc, "3.2 Architecture Flow", 2)
    num_id = create_decimal_numbering(doc)
    for item in [
        "Source systems emit telemetry or documents through signed webhooks, APIs, scheduled jobs, or manual upload.",
        "The middleware validates the payload envelope, enforces tenant and RBAC context, stores the raw evidence in S3, and routes only validated context to the proper agent.",
        "The agent returns structured JSON, not free-form prose. The middleware validates that JSON against schema and allow-list rules.",
        "Accepted decisions become graph updates, pending approval items, evidence edges, risk updates, tickets, or notifications.",
        "Every AI decision receives a cognitive audit log, hash, model/version metadata, and graph relationship to the agent that produced it.",
    ]:
        add_number(doc, item, num_id)
    add_heading(doc, "3.3 Deployment Boundary", 2)
    add_para(
        doc,
        "The UI and graph store should remain AWS-native for operational simplicity. If model orchestration runs in another cloud, engineers must use workload identity federation and short-lived credentials. Do not store long-lived AWS access keys in AI containers."
    )


def data_model(doc):
    add_heading(doc, "4. Control-Centric Graph Data Model", 1)
    add_para(
        doc,
        "The Control node is the product's nucleus. Every other major entity should either connect directly to a control or connect through a path that explains control impact. This enables the product to answer practical questions such as: what proves this control, what assets does it protect, what requirements does it satisfy, who owns it, what vendors does it rely on, and what risk exposure changes if it fails?"
    )
    add_heading(doc, "4.1 Core Node Types", 2)
    add_table(
        doc,
        ["Node", "Purpose", "Required Properties"],
        [
            ("Control", "Canonical internal control record", "tenant_id, control_id, name, description, owner_id, implementation_status, control_type, automation_level, criticality, kpi_kri, fair_parameters"),
            ("Requirement", "External framework or regulatory requirement", "tenant_id, requirement_id, framework, version, citation_id, text, jurisdiction, effective_date, status"),
            ("Asset", "System, service, cloud resource, data store, IAM platform, container, application, or process", "tenant_id, asset_id, asset_type, name, environment, criticality, data_classification, owner_id"),
            ("Evidence", "Raw or processed proof of control operation", "tenant_id, evidence_id, evidence_type, collection_method, source_system, s3_uri, s3_version_id, collected_at, valid_until, hash"),
            ("RiskScenario", "FAIR scenario or enterprise risk record", "tenant_id, scenario_id, name, scope, loss_event, asset_id, lef_parameters, lm_parameters, ale_p10, ale_p50, ale_p90"),
            ("Vendor", "Third party or service provider", "tenant_id, vendor_id, name, criticality_tier, data_access, business_owner, assessment_status, external_rating"),
            ("PolicyDocument", "Policy, SOP, standard, procedure, or generated draft", "tenant_id, document_id, title, document_type, owner_id, status, s3_uri, approved_version"),
            ("UserTeam", "Human owner, approver, team, or escalation group", "tenant_id, principal_id, name, email, roles, assigned_controls"),
            ("Agent", "AI agent identity and runtime version", "tenant_id, agent_id, agent_name, model_provider, model_version, service_account, permissions_boundary, lifecycle_status"),
            ("AuditLog", "Immutable record of human and AI decisions", "tenant_id, audit_id, actor_id, action, context_hash, output_hash, timestamp, accepted_by_middleware"),
        ],
        [1600, 3100, 4660],
    )
    add_heading(doc, "4.2 Relationship Vocabulary", 2)
    add_table(
        doc,
        ["Edge", "From", "To", "Meaning / Key Properties"],
        [
            ("IMPLEMENTED_ON", "Control", "Asset", "The control operates in a specific technical or process context. Properties: asset_context, testing_cadence, inheritance, scope_status."),
            ("SATISFIES", "Control", "Requirement", "The control maps to a framework requirement. Properties: coverage_percentage, ai_mapping_confidence, state, approved_by, approved_at."),
            ("PROVED_BY", "Control", "Evidence", "Evidence supports control operation. Properties: valid_until, telemetry_source, evidence_scope, context_hash, state."),
            ("MITIGATES", "Control", "RiskScenario", "The control reduces frequency or magnitude for a defined scenario. Properties: lef_effect, lm_effect, confidence, assumption_set_id."),
            ("OWNED_BY", "Control", "UserTeam", "Accountability and escalation. Properties: role, escalation_path, last_notified."),
            ("RELATED_TO", "Control", "Control", "Control dependency, prerequisite, overlap, compensating relationship, or conflict. Properties: relationship_type, rationale."),
            ("GOVERNS", "PolicyDocument", "Control", "Document defines or supports the control. Properties: policy_section, effective_version."),
            ("RELIED_UPON_FOR", "Vendor", "Control", "Third-party dependency supports a control. Properties: dependency_type, assessment_status."),
            ("PROCESSES_DATA_FOR", "Vendor", "Asset", "Vendor processes data for a business asset. Properties: data_type, region, processing_role."),
            ("EVALUATED_BY", "Control", "Agent", "AI agent produced a decision or recommendation. Properties: model_version, prompt_version, context_hash, verdict."),
        ],
        [1500, 1300, 1500, 5060],
    )
    add_heading(doc, "4.3 Control Node JSON Schema Skeleton", 2)
    add_code_block(
        doc,
        """{
  "$schema": "https://json-schema.org/draft/2020-12/schema",
  "title": "ControlNode",
  "type": "object",
  "required": ["tenant_id", "control_id", "name", "owner_id", "implementation_status"],
  "properties": {
    "tenant_id": {"type": "string"},
    "node_type": {"const": "Control"},
    "control_id": {"type": "string"},
    "name": {"type": "string"},
    "description": {"type": "string"},
    "control_type": {"enum": ["Preventive", "Detective", "Corrective", "Directive"]},
    "automation_level": {"enum": ["Manual", "Partially Automated", "Fully Automated", "Agentic"]},
    "owner_id": {"type": "string"},
    "implementation_status": {"enum": ["Not Started", "In Progress", "Implemented", "Degraded", "Failed", "Retired"]},
    "criticality": {"enum": ["Low", "Medium", "High", "Mission Critical"]},
    "kpi_kri": {"type": "object"},
    "fair_parameters": {"type": "object"},
    "ai_metadata": {"type": "object"}
  }
}""",
    )
    add_heading(doc, "4.4 Data Model Rules", 2)
    for item in [
        "Every node and edge must carry tenant_id. The middleware must inject tenant filters into every query, including AI-generated read-only queries.",
        "A control can be implemented differently per asset; store that difference on the IMPLEMENTED_ON edge instead of duplicating controls for every context.",
        "AI-generated mappings and high-risk changes begin as PENDING_APPROVAL, not ACTIVE.",
        "Evidence edges must store exact S3 object version identifiers so auditors retrieve the exact artifact reviewed at decision time.",
        "Never store full evidence bodies in the vector index. Store graph identifiers and minimal metadata; retrieve authoritative content from S3 and Neptune.",
    ]:
        add_bullet(doc, item)


def modules(doc):
    add_heading(doc, "5. Product Modules and Functional Requirements", 1)
    modules_rows = [
        ("Control Center", "Create, edit, assign, test, monitor, and report on controls; show connected requirements, assets, policies, evidence, risks, vendors, and AI decisions.", "Control list, control detail, relationship graph, owner workflow, status history, KPI/KRI cards."),
        ("Framework Mapper", "Import framework requirements, run semantic matching, propose SATISFIES edges, route low-confidence mappings for review, track coverage and gaps.", "Framework upload, mapping queue, requirement detail, coverage matrix, approval history."),
        ("Evidence Pipeline", "Collect, store, dedupe, validate, and reuse evidence across controls, frameworks, audits, and vendors.", "Evidence library, evidence request workflow, source system trace, validity dates, audit package export."),
        ("FAIR Risk Engine", "Translate control health into frequency and magnitude assumptions, run Monte Carlo simulations, and show percentile loss outputs.", "Scenario builder, assumption editor, simulation results, risk appetite alerts, executive dashboard."),
        ("AI Knowledge System", "Let users query the graph in natural language while enforcing read-only constraints and source-grounded answers.", "Conversational query, generated query preview for analysts, cited graph paths, answer confidence."),
        ("Third-Party Risk", "Represent vendors as graph nodes, extract vendor evidence from assessments, connect vendor posture to internal controls and assets.", "Vendor profile, assessment intake, dependency graph, monitoring alerts, vendor risk rollup."),
        ("Policy and Document Generator", "Draft policies, SOPs, standards, and audit narratives from control graph context, with analyst review before publication.", "Template picker, graph scope selector, rich-text draft editor, approval workflow, versioning."),
        ("Remediation Engine", "Recommend or trigger pre-approved remediation playbooks based on control failure, asset criticality, and policy guardrails.", "Remediation queue, playbook registry, approval gate, rollback status, post-remediation evidence."),
        ("Administration", "Manage tenants, roles, integrations, guardrails, retention, agent lifecycle, risk appetite settings, and model versions.", "Settings, integration marketplace, RBAC console, tenant configuration, agent registry."),
    ]
    add_table(doc, ["Module", "Requirements", "Key Views"], modules_rows, [1700, 4600, 3060])
    add_heading(doc, "5.1 Cross-Module Global Filters", 2)
    add_para(
        doc,
        "The UI should preserve context across modules. If a user filters the Control Center to AWS production assets and NIST CSF Protect outcomes, then framework coverage, evidence, risk, and approvals views should inherit that scope unless the user clears it. Global filters should include tenant, framework, business unit, asset type, environment, vendor tier, owner, status, criticality, data classification, and date window."
    )
    add_heading(doc, "5.2 Time-to-Value Requirements", 2)
    for item in [
        "Day 0 onboarding should connect at least one cloud source and one identity source, run discovery, create initial Asset nodes, and populate a starter Control Center.",
        "The first dashboard should show quick wins: controls automatically validated, controls missing evidence, top degraded assets, and framework requirements with likely coverage.",
        "Role-specific onboarding should guide system admins through integrations, analysts through mapping approvals, and control owners through evidence tasks.",
    ]:
        add_bullet(doc, item)


def ai_workflows(doc):
    add_heading(doc, "6. AI Agent Workflows", 1)
    add_para(
        doc,
        "Agents are product capabilities with identities, permissions, lifecycle states, audit logs, and failure modes. They should be treated like governed digital actors, not generic background jobs."
    )
    add_table(
        doc,
        ["Agent", "Trigger", "Allowed Outputs", "Human Gate"],
        [
            ("Evidence Validation Agent", "Validated telemetry, manual upload, scheduled control test", "Implemented/Degraded/Failed recommendation, PROVED_BY edge proposal, missing evidence summary", "Required for Tier 1 control failure reversal or low-confidence evidence"),
            ("Framework Mapping Agent", "Framework upload or requirement update", "Proposed SATISFIES edges with coverage and confidence, gap summary", "Required for all new mappings below high-confidence threshold and for regulated frameworks"),
            ("FAIR Risk Agent", "Control degraded/failed, asset criticality change, scenario edit, vendor status change", "Updated LEF/LM assumptions, Monte Carlo results, risk narrative, appetite alert", "Required before changing board-facing risk acceptance state"),
            ("Knowledge System Agent", "Natural language question", "Read-only answer with graph path, cited nodes/edges, confidence, query trace", "No writes allowed"),
            ("Third-Party Assessment Agent", "Vendor documentation upload or external risk feed event", "Vendor gap findings, relied-upon control impact, assessment summary", "Required for vendor tier changes and control degradation based only on unstructured docs"),
            ("Policy Drafting Agent", "User selects document type and graph scope", "Draft policy/SOP/standard mapped to controls and requirements", "Publication requires human approval"),
            ("Remediation Recommendation Agent", "Config drift, vulnerability, or failed control", "Pre-approved playbook recommendation or pending execution request", "Required for production or mission-critical assets"),
        ],
        [2200, 2500, 3100, 1560],
    )
    add_heading(doc, "6.1 Evidence Validation Workflow", 2)
    num_id = create_decimal_numbering(doc)
    for item in [
        "Retrieve the Control_Context from Neptune, including control requirements, asset-specific edge context, owner, current status, and evidence validity rules.",
        "Retrieve and normalize the Telemetry_Payload from the validated webhook or manual evidence intake.",
        "Prompt the Evidence Validation Agent using structured output and deterministic instructions.",
        "Validate the agent output against JSON schema, confidence thresholds, and allowed actions.",
        "Write the raw evidence to S3 with Object Lock and versioning, create or update the Evidence node, and create a PROVED_BY edge only if validation passes.",
        "If evidence is missing or failing, notify the control owner, create a task, and trigger the FAIR Risk Agent when material.",
    ]:
        add_number(doc, item, num_id)
    add_heading(doc, "6.2 Framework Mapping Workflow", 2)
    num_id = create_decimal_numbering(doc)
    for item in [
        "Ingest framework requirements as Requirement nodes with framework, version, citation, requirement text, and jurisdiction metadata.",
        "Embed the requirement text and retrieve top candidate Control IDs from the vector index.",
        "Hydrate only those candidate controls from Neptune and pass them to the Framework Mapping Agent.",
        "Create proposed SATISFIES edges in PENDING_APPROVAL with coverage, confidence, and reasoning.",
        "Route proposals to the Pending Approvals Queue with a side-by-side requirement/control review surface.",
        "On approval, activate the edge and update framework coverage. On rejection, preserve the rejection rationale for future prompt and retrieval tuning.",
    ]:
        add_number(doc, item, num_id)
    add_heading(doc, "6.3 Prompt Contract Pattern", 2)
    add_code_block(
        doc,
        """System prompt pattern:
ROLE: Define the agent's narrow professional role.
INSTRUCTIONS: Define the exact input objects and what comparison or analysis is allowed.
STEPS: Force a repeatable sequence: parse context, parse payload, evaluate, produce output.
EXPECTATIONS: Forbid hallucinated controls, unsupported assumptions, compensating controls not provided in context, and free-form prose outside JSON.
OUTPUT: Require strict JSON with IDs, verdict, confidence, reasoning, missing evidence, proposed action, and state."""
    )


def risk_engine(doc):
    add_heading(doc, "7. FAIR Risk and Quantitative Control Health", 1)
    add_para(
        doc,
        "Risk should be modeled as scenario-specific, probabilistic, and control-sensitive. The platform should not convert FAIR into a hidden score. It should expose assumptions, ranges, distributions, confidence, and the controls that influence loss event frequency or loss magnitude."
    )
    add_heading(doc, "7.1 FAIR Data Model", 2)
    add_table(
        doc,
        ["FAIR Concept", "Platform Representation", "Data Sources"],
        [
            ("Loss Event Frequency", "Scenario property derived from threat frequency and susceptibility/resistance strength", "Threat intel, SIEM, vulnerability scanners, historical incidents, control performance"),
            ("Loss Magnitude", "Range distribution across primary and secondary loss categories", "Finance inputs, incident response cost models, legal/regulatory estimates, downtime estimates, customer/reputation assumptions"),
            ("Control Effect", "MITIGATES edge properties that describe how a control changes frequency or magnitude", "Control owner input, historical effectiveness, AI-assisted scenario analysis, analyst approval"),
            ("Uncertainty", "Min/most-likely/max or distribution parameters", "Subject matter estimates, telemetry ranges, sensitivity analysis"),
            ("Output", "ALE percentiles and narrative explanation", "Monte Carlo simulation microservice"),
        ],
        [2100, 4100, 3160],
    )
    add_heading(doc, "7.2 Monte Carlo Simulation Requirements", 2)
    for item in [
        "Support at least 10,000 simulation iterations per scenario, configurable by tenant and scenario materiality.",
        "Return P10, P50, P90, expected value, probability of exceeding appetite, and sensitivity drivers.",
        "Store the exact assumption set used for each run so results are reproducible.",
        "Trigger re-runs when material control health, vendor status, vulnerability state, asset criticality, or threat telemetry changes.",
        "Use the AI agent for scenario narrative, assumption translation, and stakeholder explanation, not as the mathematical simulation engine.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.3 Risk Dashboard Requirements", 2)
    add_para(
        doc,
        "The executive view should show top scenarios by P90 exposure, control failures with material effects, trend in aggregate control health, risk appetite breaches, and the graph path from scenario to control to evidence. Analyst views should expose assumptions and allow controlled edits with approval history."
    )


def api_integrations(doc):
    add_heading(doc, "8. API, Evidence, and Integration Architecture", 1)
    add_heading(doc, "8.1 Common Telemetry Envelope", 2)
    add_code_block(
        doc,
        """{
  "event_id": "uuid",
  "timestamp": "2026-07-30T12:00:00Z",
  "tenant_id": "tenant_123",
  "source_system": "AWS_SecurityHub",
  "event_type": "Config_Drift",
  "signature_version": "v1",
  "payload": {}
}""",
    )
    add_para(
        doc,
        "Every integration should normalize into a common envelope before agent evaluation. The API must verify HMAC signatures, reject stale timestamps, enforce idempotency on event_id, validate payload schema, and store the raw event before producing graph mutations."
    )
    add_heading(doc, "8.2 Integration Prioritization", 2)
    add_table(
        doc,
        ["Priority", "Integration Category", "Why It Comes First", "Initial Targets"],
        [
            ("1", "Cloud posture and asset discovery", "Creates the asset graph and high-volume evidence baseline", "AWS Security Hub, AWS Config, CloudTrail, IAM Access Analyzer"),
            ("2", "Identity and access management", "Access controls are heavily audited and tie to many frameworks", "Okta or Microsoft Entra ID, AWS IAM logs"),
            ("3", "Vulnerability and threat telemetry", "Feeds control degradation and FAIR frequency/susceptibility assumptions", "Inspector, Qualys/Tenable/Rapid7, SIEM/XDR"),
            ("4", "Workflow systems", "Pushes remediation into existing operational habits", "Jira, ServiceNow, Slack/Teams notifications"),
            ("5", "Third-party risk and legacy imports", "Bootstraps vendor and risk context without manual re-entry", "OneTrust export, ProcessUnity/CyberGRX-style risk feeds, vendor document upload"),
            ("6", "Document repositories and policy systems", "Links policies, SOPs, and standards to controls", "SharePoint/Drive/Confluence or existing document store"),
        ],
        [900, 2500, 3300, 2660],
    )
    add_heading(doc, "8.3 Immutable Evidence Pipeline", 2)
    for item in [
        "Write every raw webhook, upload, AI context log, and final decision artifact to an S3 bucket with versioning enabled.",
        "Apply Object Lock retention policies to evidence and AI decision logs according to legal and audit requirements.",
        "Store the bucket, key, version ID, content hash, collected_at, valid_until, and retention mode on the Evidence node or related edge.",
        "When a user clicks evidence in the UI, retrieve the exact object version referenced by the graph edge.",
        "If evidence must be annotated later, create a new object version or annotation artifact; do not mutate the original locked evidence.",
    ]:
        add_bullet(doc, item)


def ui_rbac(doc):
    add_heading(doc, "9. User Experience and RBAC", 1)
    add_heading(doc, "9.1 Role-Specific Views", 2)
    add_table(
        doc,
        ["Persona", "Primary Jobs", "Default Landing View"],
        [
            ("System Admin", "Configure tenants, integrations, RBAC, retention, guardrails, and agent lifecycle", "Admin health and integration status"),
            ("Compliance Analyst", "Manage controls, approve mappings, prepare audits, review evidence, tune frameworks", "Control Center and pending approvals"),
            ("Control Owner", "Respond to evidence requests, fix degraded controls, confirm remediation", "Assigned control tasks and daily digest"),
            ("Third-Party Risk Analyst", "Manage vendor intake, review vendor evidence, assess vendor control dependencies", "Vendor risk queue"),
            ("Executive / Board", "Understand material exposure, control health, risk trends, and compliance readiness", "Global FAIR and control health dashboard"),
            ("Auditor", "Review evidence, control mappings, audit trails, and immutable proof packages", "Audit package and evidence trail"),
        ],
        [2100, 5000, 2260],
    )
    add_heading(doc, "9.2 Required UI Views", 2)
    for item in [
        "Global Risk Dashboard: FAIR P10/P50/P90 exposure, appetite breaches, control health rollup, top scenarios, and material control failures.",
        "Control Detail: metadata, owner, implementation status, KPIs/KRIs, sub-graph, evidence, requirements, assets, policies, risk scenarios, vendors, AI history.",
        "Graph Explorer: interactive subgraph with filters, path explanations, and edge detail side panels.",
        "Pending Approvals Queue: AI-proposed mappings, evidence verdicts, remediation requests, confidence, diff viewer, approve/reject with rationale.",
        "Evidence Library: search, validity status, source system, linked controls, immutable object version, audit package inclusion.",
        "Framework Coverage: requirement list, mapped controls, coverage percentage, gaps, version changes, approval history.",
        "Integration Marketplace: OAuth/API connection flows, health checks, last sync, permission scopes, connector logs.",
        "Settings: tenant scoping, risk appetite, FAIR defaults, workflow templates, notification rules, retention, model versions.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.3 RBAC Matrix", 2)
    add_table(
        doc,
        ["Action", "Admin", "Analyst", "Owner", "TPRM", "Executive", "Auditor"],
        [
            ("View global FAIR metrics", "Yes", "Yes", "No", "No", "Yes", "Scoped"),
            ("View control graph", "Yes", "Yes", "Assigned", "Assigned vendor", "Read-only", "Scoped read-only"),
            ("Edit control metadata", "Yes", "Yes", "Assigned fields", "No", "No", "No"),
            ("Approve AI mappings", "Yes", "Yes", "No", "Vendor scope only", "No", "No"),
            ("Upload manual evidence", "Yes", "Yes", "Assigned", "Vendor scope", "No", "No"),
            ("Manage vendors", "Yes", "Yes", "No", "Yes", "Read-only", "Scoped read-only"),
            ("Trigger Monte Carlo", "Yes", "Yes", "No", "No", "No", "No"),
            ("Modify guardrails/settings", "Yes", "No", "No", "No", "No", "No"),
        ],
        [2500, 1000, 1000, 1200, 1100, 1200, 1360],
    )
    add_heading(doc, "9.4 Trust UX", 2)
    add_para(
        doc,
        "Every AI-assisted decision should show confidence, evidence source, raw context, graph path, generated reasoning, middleware validation result, and human approval state. For high-risk actions, the UI should introduce deliberate friction: typed confirmation, second approver, or change-window enforcement."
    )


def security_ops(doc):
    add_heading(doc, "10. Security, Privacy, Guardrails, and Operations", 1)
    add_heading(doc, "10.1 AI Write Guardrails", 2)
    add_table(
        doc,
        ["Guardrail", "Requirement"],
        [
            ("No direct DB access", "Agents cannot connect to Neptune. They return structured JSON to middleware only."),
            ("Agent-specific identities", "Each agent has a service account and mapped IAM role with least-privilege permissions."),
            ("Action allow-list", "Middleware permits only approved actions per agent, such as CREATE_EDGE:PROVED_BY for Evidence Agent."),
            ("Global deny-list", "DELETE_NODE:CONTROL, unbounded updates, schema-altering mutations, and tenant changes are rejected."),
            ("Schema validation", "Every AI output must validate before query construction."),
            ("Traversal limits", "Middleware-generated queries must target explicit IDs and enforce maximum hop counts."),
            ("Human-in-the-loop", "High-risk writes and low-confidence mappings remain PENDING_APPROVAL."),
            ("Immutable audit", "Prompt, context, output, action verdict, and hash are stored in a locked evidence log."),
        ],
        [2600, 6760],
    )
    add_heading(doc, "10.2 Data Isolation and Residency", 2)
    for item in [
        "Every data object must include tenant_id and business-unit scope where applicable.",
        "Middleware must inject tenant_id into every query and reject user-supplied tenant overrides.",
        "Evidence buckets should be partitioned by tenant and environment, with KMS keys and bucket policies scoped accordingly.",
        "Data residency must be configurable for regulated customers; model processing regions need explicit approval.",
        "PII, secrets, and sensitive evidence should be redacted or tokenized before AI context injection unless explicitly required.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.3 Agent Lifecycle Monitoring", 2)
    for item in [
        "Track tokens, latency, tool-call rate, graph traversal depth, rejection rate, approval rate, and hallucination/schema-failure rate by agent and model version.",
        "Implement circuit breakers for recursive query generation, abnormal read volume, repeated schema failures, or unauthorized action attempts.",
        "Maintain Agent nodes with lifecycle_status values such as ACTIVE, CANARY, DEPRECATED, RETIRED, and DISABLED.",
        "When a model version is retired or found defective, query EVALUATED_BY edges to identify affected controls and route them for re-review.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.4 BCP/DR and Operational Readiness", 2)
    add_table(
        doc,
        ["Area", "Requirement"],
        [
            ("Recovery objectives", "Define RPO and RTO for Neptune, evidence store, orchestration hub, and UI before production."),
            ("Fallback mode", "If KMS Text-to-Graph fails repeatedly, degrade to guided search and open a support ticket with failed prompts attached."),
            ("Control-owner enablement", "Provide daily digest and direct task links instead of forcing owners to live inside the dashboard."),
            ("SLA model", "Tier 1 control failure should page appropriate owners quickly; Tier 3 remediation may be ticket-based or automated."),
            ("Day 2 runbook", "Document support ownership, connector failures, evidence sync failures, model rollback, and graph repair procedures."),
        ],
        [2300, 7060],
    )


def testing_backlog(doc):
    add_heading(doc, "11. Delivery Plan and Acceptance Criteria", 1)
    add_heading(doc, "11.1 Phased Delivery", 2)
    add_table(
        doc,
        ["Phase", "Goal", "Primary Deliverables", "Exit Criteria"],
        [
            ("0. Discovery and Charter", "Lock product assumptions, scope, personas, primary framework, and security posture", "Charter, PRD, architecture decision records, threat model, data classification", "Stakeholders approve scope and non-goals"),
            ("1. Control Center Foundation", "Create the graph-backed system of record", "Control/Asset/Requirement/Evidence schemas, UI CRUD, owner assignment, basic graph view", "Users can create controls and connect them to assets and requirements"),
            ("2. Evidence and First Integrations", "Prove controls continuously", "Telemetry envelope, AWS integration, IAM integration, S3 immutable evidence, Evidence Agent", "Evidence validates a control and creates a PROVED_BY edge with raw artifact"),
            ("3. Framework Mapping and Approvals", "Reduce duplicate compliance work", "Framework import, vector retrieval, Mapping Agent, approval queue, coverage dashboard", "Analysts approve/reject proposed mappings and coverage updates accurately"),
            ("4. FAIR Risk Engine", "Translate control health into business exposure", "Scenario model, Monte Carlo service, risk dashboard, appetite alerts", "A degraded control updates scenario exposure with reproducible assumptions"),
            ("5. TPRM and Policy Workflows", "Connect vendors and documentation to controls", "Vendor graph, assessment agent, policy generator, document versioning", "Vendor evidence and policies map to controls with review history"),
            ("6. Remediation and Scale Hardening", "Close the loop safely", "Playbook registry, remediation recommendations, HITL execution, rollback evidence, agent monitoring", "Approved playbooks remediate low-risk drift and preserve audit trail"),
        ],
        [1500, 2500, 3200, 2160],
    )
    add_heading(doc, "11.2 Product Backlog", 2)
    add_table(
        doc,
        ["Epic", "Representative User Stories"],
        [
            ("Control Center", "As an analyst, I can create a control with owner, status, automation level, KPI/KRI, and framework relevance. As a control owner, I can see only controls assigned to me."),
            ("Graph Relationships", "As an analyst, I can connect a control to assets, requirements, evidence, vendors, policies, and risk scenarios. As an auditor, I can inspect the full evidence path."),
            ("Evidence Automation", "As a system, I can ingest AWS/IAM telemetry, validate it, store immutable raw evidence, and update the control graph."),
            ("AI Approvals", "As an analyst, I can review AI proposals side by side with source context, confidence, and reasoning, then approve or reject with rationale."),
            ("FAIR Risk", "As a risk manager, I can define a scenario, attach controls, run Monte Carlo simulation, and see how degraded controls affect P10/P50/P90 loss."),
            ("Knowledge Query", "As an executive, I can ask what controls are failing for a critical asset and receive a read-only answer tied to graph paths."),
            ("Vendor Risk", "As a TPRM analyst, I can upload a SOC 2 report, extract control coverage, and see internal controls that rely on the vendor."),
            ("Operations", "As an admin, I can monitor agent behavior, disable an agent, rotate credentials, inspect connector health, and export audit logs."),
        ],
        [2200, 7160],
    )
    add_heading(doc, "11.3 Testing Strategy", 2)
    for item in [
        "Unit and contract tests for JSON schemas, webhook validation, idempotency, tenant scoping, and graph query construction.",
        "Graph integrity tests for required node properties, allowed edge types, orphaned nodes, tenant leaks, and invalid relationship states.",
        "AI evaluation tests with golden datasets for evidence validation, mapping precision/recall, prompt injection resistance, and refusal behavior.",
        "Adversarial red-team tests that attempt prompt injection, unauthorized graph mutation, tenant data leakage, fake evidence, replay attacks, and overbroad remediation.",
        "Parallel-run UAT against the current risk register and evidence process before production cutover.",
        "Operational chaos tests for connector outages, model failure, vector index staleness, S3 write failures, and Neptune failover.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "11.4 Definition of Done", 2)
    add_para(
        doc,
        "A feature is not done when it renders in the UI. It is done when it enforces tenant isolation, writes complete audit records, has schema and RBAC tests, exposes error states to users, preserves immutable evidence where applicable, and can be explained through the control graph."
    )


def appendices(doc):
    add_heading(doc, "12. Open Decisions", 1)
    for item in [
        "Primary framework for MVP: NIST CSF 2.0 is recommended, but final selection should match the first target customer or internal compliance need.",
        "Model provider and hosting boundary: confirm whether orchestration is AWS-only or cross-cloud, because identity federation and residency requirements depend on that decision.",
        "Commercial SaaS versus internal enterprise tool: multi-tenancy, pricing, customer onboarding, trust center, and marketplace depth depend on the answer.",
        "Risk appetite governance: determine who can set financial thresholds and who can approve acceptance of material risk.",
        "Evidence retention period: align S3 Object Lock settings with legal, audit, and customer obligations before production.",
        "Automated remediation policy: define asset tiers and control categories where autonomous playbooks are allowed.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "13. Source-Informed Reference Notes", 1)
    refs = [
        ("Drata controls and evidence", "Control ownership, cross-framework mapping, centralized/reusable evidence, and continuous visibility are baseline buyer expectations.", "https://drata.com/products/compliance/controls-and-evidence"),
        ("Vanta automated compliance", "Continuous monitoring, automated evidence, alerts, remediation support, and stakeholder collaboration are common compliance automation expectations.", "https://www.vanta.com/products/automated-compliance"),
        ("Archer IT and Security Risk Management", "Enterprise GRC buyers expect controls, vulnerabilities, audit findings, obligations, reporting, policy management, continuous controls monitoring, and asset context.", "https://www.archerirm.com/it-security-risk-management"),
        ("OneTrust GRC overview", "Scalable GRC commonly includes automated risk assessments, policy management, compliance tracking, and unified governance/risk/operational data.", "https://www.onetrust.com/glossary/governance-risk-and-compliance-grc/"),
        ("Compyl framework/evidence positioning", "One control library, many framework mappings, evidence reuse, automated evidence, and real-time program views should be considered product table stakes.", "https://compyl.com/m-frameworks/"),
        ("ProcessUnity / CyberGRX risk exchange", "Third-party risk is strengthened by reusable assessments, external risk intelligence, and updated vendor posture views.", "https://www.processunity.com/resources/blogs/facilitate-third-party-risk-assessments-using-global-risk-exchange/"),
        ("NIST CSF 2.0", "The framework supports flexible outcomes, profiles, tiers, governance, supply chain emphasis, and continuous cybersecurity risk management.", "https://nvlpubs.nist.gov/nistpubs/CSWP/NIST.CSWP.29.pdf"),
        ("FAIR Standard v3.0", "FAIR frames risk as probable frequency and magnitude of future loss, derived from Loss Event Frequency and Loss Magnitude.", "https://www.fairinstitute.org/hubfs/Standards%20Artifacts/Factor%20Analysis%20of%20Information%20Risk%20%28FAIR%29%20Standard%20v3.0%20%28January%202025%29.pdf"),
        ("Amazon Neptune data model", "Neptune supports graph relationships and properties, suitable for explicit control-to-entity relationship traversal.", "https://docs.aws.amazon.com/neptune/latest/userguide/feature-overview-data-model.html"),
        ("Amazon S3 Object Lock", "Object Lock supports WORM-style retention to help prevent evidence objects from deletion or overwrite for fixed periods or indefinitely.", "https://docs.aws.amazon.com/AmazonS3/latest/userguide/object-lock.html"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [2300, 4560, 2500])
    for i, h in enumerate(["Reference", "How It Informed the Plan", "URL"]):
        shade_cell(table.rows[0].cells[i], LIGHT_GRAY)
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(run, size=9.5, bold=True)
    for ref, note, url in refs:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(ref)
        set_run_font(cells[0].paragraphs[0].runs[0], size=9.5)
        cells[1].paragraphs[0].add_run(note)
        set_run_font(cells[1].paragraphs[0].runs[0], size=9.5)
        add_hyperlink(cells[2].paragraphs[0], url, url)
    set_table_width(table, [2300, 4560, 2500])
    add_heading(doc, "14. Glossary", 1)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ("Control Center", "The control-backed system of record and primary workspace for the platform."),
            ("Control graph", "The graph representation of controls and all connected entities, including assets, evidence, requirements, vendors, risk scenarios, policies, users, and agents."),
            ("PROVED_BY", "The graph edge linking a control to evidence that supports its operating effectiveness."),
            ("SATISFIES", "The graph edge linking an internal control to an external framework requirement."),
            ("PENDING_APPROVAL", "State for AI-proposed or high-risk graph changes awaiting human review."),
            ("Cognitive audit log", "The exact prompt, context, model version, output, validation result, and hash for an AI decision."),
            ("GraphRAG / Text-to-Graph", "A natural language query workflow that uses graph schema and graph traversals rather than loose document search alone."),
            ("FAIR", "Factor Analysis of Information Risk, used here for quantitative scenario-based cyber risk analysis."),
        ],
        [2200, 7160],
    )


def main():
    doc = setup_doc()
    title_page(doc)
    market_context(doc)
    charter_scope(doc)
    architecture(doc)
    data_model(doc)
    modules(doc)
    ai_workflows(doc)
    risk_engine(doc)
    api_integrations(doc)
    ui_rbac(doc)
    security_ops(doc)
    testing_backlog(doc)
    appendices(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
